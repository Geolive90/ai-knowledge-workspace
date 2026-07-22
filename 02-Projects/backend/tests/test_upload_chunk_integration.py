"""Integration tests for upload with chunk persistence."""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from app.dependencies import get_db
from app.main import app
from app.models.document import Document
from app.models.document_chunk import DocumentChunk
from app.utils import file_handler


@pytest.fixture
def upload_folder(tmp_path, monkeypatch):
    upload_dir = tmp_path / "uploads"
    upload_dir.mkdir()
    monkeypatch.setattr(file_handler, "UPLOAD_FOLDER", upload_dir)
    return upload_dir


@pytest.fixture
def client(db_session, upload_folder):
    def override_get_db():
        try:
            yield db_session
        finally:
            pass

    app.dependency_overrides[get_db] = override_get_db

    with TestClient(app) as test_client:
        yield test_client

    app.dependency_overrides.clear()


def register_and_login(client: TestClient, email: str, password: str) -> str:
    client.post(
        "/auth/register",
        json={
            "email": email,
            "full_name": "Upload Test User",
            "password": password,
        },
    )
    login_response = client.post(
        "/auth/login",
        data={"username": email, "password": password},
    )
    assert login_response.status_code == 200
    return login_response.json()["access_token"]


def auth_headers(token: str) -> dict[str, str]:
    return {"Authorization": f"Bearer {token}"}


def test_authenticated_txt_upload_persists_document_and_chunks(
    client,
    db_session,
) -> None:
    token = register_and_login(client, "upload-txt@example.com", "password123")
    content = ("Upload integration paragraph.\n\n" + "word " * 250).encode("utf-8")

    response = client.post(
        "/documents/upload",
        headers=auth_headers(token),
        files={"file": ("integration.txt", BytesIO(content), "text/plain")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert set(payload.keys()) == {
        "message",
        "document_id",
        "filename",
        "extracted_character_count",
        "text_preview",
    }
    assert payload["filename"] == "integration.txt"
    assert payload["message"] == "Document uploaded and text extracted successfully."

    document = (
        db_session.query(Document)
        .filter(Document.id == payload["document_id"])
        .first()
    )
    assert document is not None
    assert document.extracted_text is not None
    assert len(document.chunks) > 0
    assert payload["extracted_character_count"] == len(document.extracted_text)
    assert payload["text_preview"] == document.extracted_text[:300]

    for chunk in document.chunks:
        assert (
            document.extracted_text[
                chunk.character_start : chunk.character_end
            ]
            == chunk.chunk_text
        )


def test_crlf_upload_response_uses_persisted_normalized_text(
    client,
    db_session,
) -> None:
    token = register_and_login(client, "upload-crlf@example.com", "password123")
    raw_text = "Alpha\r\n\r\n" + ("beta " * 120)
    content = raw_text.encode("utf-8")

    response = client.post(
        "/documents/upload",
        headers=auth_headers(token),
        files={"file": ("crlf.txt", BytesIO(content), "text/plain")},
    )

    assert response.status_code == 200
    payload = response.json()

    document = (
        db_session.query(Document)
        .filter(Document.id == payload["document_id"])
        .first()
    )
    assert document is not None
    assert document.extracted_text is not None
    assert "\r" not in document.extracted_text
    assert payload["extracted_character_count"] == len(document.extracted_text)
    assert payload["text_preview"] == document.extracted_text[:300]


def test_list_retrieve_download_delete_still_work(
    client,
    db_session,
    upload_folder,
) -> None:
    token = register_and_login(client, "upload-crud@example.com", "password123")
    content = b"CRUD verification content for upload integration."

    upload_response = client.post(
        "/documents/upload",
        headers=auth_headers(token),
        files={"file": ("crud.txt", BytesIO(content), "text/plain")},
    )
    assert upload_response.status_code == 200
    document_id = upload_response.json()["document_id"]

    list_response = client.get("/documents", headers=auth_headers(token))
    assert list_response.status_code == 200
    assert any(item["id"] == document_id for item in list_response.json())

    get_response = client.get(
        f"/documents/{document_id}",
        headers=auth_headers(token),
    )
    assert get_response.status_code == 200
    assert get_response.json()["filename"] == "crud.txt"

    download_response = client.get(
        f"/documents/{document_id}/download",
        headers=auth_headers(token),
    )
    assert download_response.status_code == 200
    assert download_response.content == content

    delete_response = client.delete(
        f"/documents/{document_id}",
        headers=auth_headers(token),
    )
    assert delete_response.status_code == 200
    assert (
        db_session.query(DocumentChunk)
        .filter(DocumentChunk.document_id == document_id)
        .count()
        == 0
    )


def test_cross_user_access_returns_404(client, db_session) -> None:
    owner_token = register_and_login(
        client,
        "owner-upload@example.com",
        "password123",
    )
    other_token = register_and_login(
        client,
        "other-upload@example.com",
        "password456",
    )

    upload_response = client.post(
        "/documents/upload",
        headers=auth_headers(owner_token),
        files={
            "file": (
                "private.txt",
                BytesIO(b"Owner-only content"),
                "text/plain",
            )
        },
    )
    document_id = upload_response.json()["document_id"]

    for method, url in [
        ("get", f"/documents/{document_id}"),
        ("get", f"/documents/{document_id}/download"),
        ("delete", f"/documents/{document_id}"),
    ]:
        response = getattr(client, method)(
            url,
            headers=auth_headers(other_token),
        )
        assert response.status_code == 404
        assert response.json()["detail"] == "Document not found."


def test_docx_upload_path_persists_document_and_chunks(
    client,
    db_session,
    tmp_path,
) -> None:
    token = register_and_login(client, "upload-docx@example.com", "password123")

    docx_path = tmp_path / "sample.docx"
    docx_document = DocxDocument()
    docx_document.add_paragraph("DOCX integration paragraph.")
    docx_document.add_paragraph("Second paragraph for chunking.")
    docx_document.save(docx_path)

    with docx_path.open("rb") as docx_file:
        response = client.post(
            "/documents/upload",
            headers=auth_headers(token),
            files={
                "file": (
                    "sample.docx",
                    docx_file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    document = (
        db_session.query(Document)
        .filter(Document.id == response.json()["document_id"])
        .first()
    )
    assert document is not None
    assert "DOCX integration paragraph." in (document.extracted_text or "")
    assert len(document.chunks) >= 1


def test_legacy_document_without_chunks_is_accessible(
    client,
    db_session,
    test_user,
) -> None:
    legacy_document = Document(
        user_id=test_user.id,
        filename="legacy-no-chunks.txt",
        file_path="legacy-no-chunks.txt",
        extracted_text="Legacy row without chunk persistence.",
    )
    db_session.add(legacy_document)
    db_session.commit()
    db_session.refresh(legacy_document)

    login_response = client.post(
        "/auth/login",
        data={
            "username": "phase3-test@example.com",
            "password": "testpassword",
        },
    )
    assert login_response.status_code == 200
    token = login_response.json()["access_token"]

    response = client.get(
        f"/documents/{legacy_document.id}",
        headers=auth_headers(token),
    )

    assert response.status_code == 200
    assert response.json()["filename"] == "legacy-no-chunks.txt"
    assert (
        db_session.query(DocumentChunk)
        .filter(DocumentChunk.document_id == legacy_document.id)
        .count()
        == 0
    )
